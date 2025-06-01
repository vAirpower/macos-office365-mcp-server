"""
Word Controller for Office 365 MCP Server
Handles all Word automation operations on macOS.
"""

import asyncio
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import logging

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from integrations.applescript_bridge import AppleScriptBridge
from utils.logger import setup_logger

logger = setup_logger(__name__)

class WordController:
    """Controller for Word operations using both AppleScript and python-docx."""
    
    def __init__(self):
        self.applescript = AppleScriptBridge()
        self.active_documents: Dict[str, Dict[str, Any]] = {}
        self.temp_dir = Path.home() / "tmp" / "office365_mcp"
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    async def create_document(
        self,
        title: str = "New Document",
        template_path: Optional[str] = None
    ) -> Dict[str, Any]:
        """Create a new Word document.
        
        Args:
            title: Document title
            template_path: Optional template file path
            
        Returns:
            Dict with document metadata
        """
        try:
            document_id = str(uuid.uuid4())
            
            # Create document using python-docx
            if template_path and Path(template_path).exists():
                doc = Document(template_path)
            else:
                doc = Document()
            
            # Add title as heading
            if title and title != "New Document":
                doc.add_heading(title, level=1)
            
            # Save temporary file
            temp_file = self.temp_dir / f"{document_id}.docx"
            doc.save(str(temp_file))
            
            # Try to open in Word via AppleScript
            applescript_success = False
            try:
                await self.applescript.open_word_file(str(temp_file))
                applescript_success = True
            except Exception as e:
                logger.warning(f"Could not open in Word app: {e}")
            
            # Store document metadata
            document_data = {
                "document_id": document_id,
                "title": title,
                "file_path": str(temp_file),
                "paragraph_count": len(doc.paragraphs),
                "applescript_available": applescript_success,
                "created_at": asyncio.get_event_loop().time()
            }
            
            self.active_documents[document_id] = {
                "metadata": document_data,
                "docx_object": doc,
                "elements": {}
            }
            
            logger.info(f"Created document: {title} ({document_id})")
            return document_data
            
        except Exception as e:
            logger.error(f"Failed to create document: {e}")
            raise
    
    async def add_heading(
        self,
        document_id: str,
        text: str,
        level: int = 1,
        style: Optional[str] = None
    ) -> Dict[str, Any]:
        """Add a heading to a document.
        
        Args:
            document_id: ID of the document
            text: Heading text
            level: Heading level (1-6)
            style: Optional style name
            
        Returns:
            Dict with operation status
        """
        try:
            if document_id not in self.active_documents:
                raise ValueError(f"Document {document_id} not found")
            
            doc = self.active_documents[document_id]["docx_object"]
            element_id = str(uuid.uuid4())
            
            # Validate level
            level = max(1, min(6, level))
            
            # Add heading
            heading = doc.add_heading(text, level=level)
            
            # Apply custom style if provided
            if style:
                try:
                    heading.style = style
                except Exception as e:
                    logger.warning(f"Could not apply style '{style}': {e}")
            
            # Store element metadata
            element_data = {
                "element_id": element_id,
                "document_id": document_id,
                "type": "heading",
                "text": text,
                "level": level,
                "style": style,
                "created_at": asyncio.get_event_loop().time()
            }
            
            self.active_documents[document_id]["elements"][element_id] = {
                "metadata": element_data,
                "element_object": heading
            }
            
            # Update document metadata
            self.active_documents[document_id]["metadata"]["paragraph_count"] = len(doc.paragraphs)
            
            # Save updated document
            temp_file = self.active_documents[document_id]["metadata"]["file_path"]
            doc.save(temp_file)
            
            logger.info(f"Added heading to document {document_id}")
            return element_data
            
        except Exception as e:
            logger.error(f"Failed to add heading: {e}")
            raise
    
    async def add_paragraph(
        self,
        document_id: str,
        text: str,
        style: Optional[str] = None,
        formatting: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Add a paragraph to a document.
        
        Args:
            document_id: ID of the document
            text: Paragraph text
            style: Optional style name
            formatting: Text formatting options
            
        Returns:
            Dict with operation status
        """
        try:
            if document_id not in self.active_documents:
                raise ValueError(f"Document {document_id} not found")
            
            doc = self.active_documents[document_id]["docx_object"]
            element_id = str(uuid.uuid4())
            
            # Add paragraph
            paragraph = doc.add_paragraph(text)
            
            # Apply style if provided
            if style:
                try:
                    paragraph.style = style
                except Exception as e:
                    logger.warning(f"Could not apply style '{style}': {e}")
            
            # Apply formatting if provided
            if formatting:
                await self._apply_paragraph_formatting(paragraph, formatting)
            
            # Store element metadata
            element_data = {
                "element_id": element_id,
                "document_id": document_id,
                "type": "paragraph",
                "text": text,
                "style": style,
                "formatting": formatting,
                "created_at": asyncio.get_event_loop().time()
            }
            
            self.active_documents[document_id]["elements"][element_id] = {
                "metadata": element_data,
                "element_object": paragraph
            }
            
            # Update document metadata
            self.active_documents[document_id]["metadata"]["paragraph_count"] = len(doc.paragraphs)
            
            # Save updated document
            temp_file = self.active_documents[document_id]["metadata"]["file_path"]
            doc.save(temp_file)
            
            logger.info(f"Added paragraph to document {document_id}")
            return element_data
            
        except Exception as e:
            logger.error(f"Failed to add paragraph: {e}")
            raise
    
    async def add_list(
        self,
        document_id: str,
        items: List[str],
        list_type: str = "bullet",
        style: Optional[str] = None
    ) -> Dict[str, Any]:
        """Add a list to a document.
        
        Args:
            document_id: ID of the document
            items: List items
            list_type: Type of list (bullet, number)
            style: Optional style name
            
        Returns:
            Dict with operation status
        """
        try:
            if document_id not in self.active_documents:
                raise ValueError(f"Document {document_id} not found")
            
            doc = self.active_documents[document_id]["docx_object"]
            element_id = str(uuid.uuid4())
            
            # Add list items
            list_elements = []
            for item in items:
                if list_type == "number":
                    paragraph = doc.add_paragraph(item, style='List Number')
                else:
                    paragraph = doc.add_paragraph(item, style='List Bullet')
                
                list_elements.append(paragraph)
            
            # Store element metadata
            element_data = {
                "element_id": element_id,
                "document_id": document_id,
                "type": "list",
                "items": items,
                "list_type": list_type,
                "style": style,
                "item_count": len(items),
                "created_at": asyncio.get_event_loop().time()
            }
            
            self.active_documents[document_id]["elements"][element_id] = {
                "metadata": element_data,
                "element_objects": list_elements
            }
            
            # Update document metadata
            self.active_documents[document_id]["metadata"]["paragraph_count"] = len(doc.paragraphs)
            
            # Save updated document
            temp_file = self.active_documents[document_id]["metadata"]["file_path"]
            doc.save(temp_file)
            
            logger.info(f"Added list to document {document_id}")
            return element_data
            
        except Exception as e:
            logger.error(f"Failed to add list: {e}")
            raise
    
    async def add_table(
        self,
        document_id: str,
        rows: int,
        columns: int,
        data: Optional[List[List[str]]] = None,
        style: Optional[str] = None
    ) -> Dict[str, Any]:
        """Add a table to a document.
        
        Args:
            document_id: ID of the document
            rows: Number of rows
            columns: Number of columns
            data: Optional table data
            style: Optional table style
            
        Returns:
            Dict with operation status
        """
        try:
            if document_id not in self.active_documents:
                raise ValueError(f"Document {document_id} not found")
            
            doc = self.active_documents[document_id]["docx_object"]
            element_id = str(uuid.uuid4())
            
            # Create table
            table = doc.add_table(rows=rows, cols=columns)
            
            # Apply style if provided
            if style:
                try:
                    table.style = style
                except Exception as e:
                    logger.warning(f"Could not apply table style '{style}': {e}")
            
            # Fill table with data if provided
            if data:
                for row_idx, row_data in enumerate(data[:rows]):
                    for col_idx, cell_data in enumerate(row_data[:columns]):
                        if row_idx < len(table.rows) and col_idx < len(table.columns):
                            table.cell(row_idx, col_idx).text = str(cell_data)
            
            # Store element metadata
            element_data = {
                "element_id": element_id,
                "document_id": document_id,
                "type": "table",
                "rows": rows,
                "columns": columns,
                "style": style,
                "has_data": bool(data),
                "created_at": asyncio.get_event_loop().time()
            }
            
            self.active_documents[document_id]["elements"][element_id] = {
                "metadata": element_data,
                "element_object": table
            }
            
            # Save updated document
            temp_file = self.active_documents[document_id]["metadata"]["file_path"]
            doc.save(temp_file)
            
            logger.info(f"Added table to document {document_id}")
            return element_data
            
        except Exception as e:
            logger.error(f"Failed to add table: {e}")
            raise
    
    async def save_document(
        self,
        document_id: str,
        file_path: str,
        format: str = "docx"
    ) -> Dict[str, Any]:
        """Save a document to file.
        
        Args:
            document_id: ID of the document
            file_path: Path to save the file
            format: File format (docx, pdf, etc.)
            
        Returns:
            Dict with operation status
        """
        try:
            if document_id not in self.active_documents:
                raise ValueError(f"Document {document_id} not found")
            
            doc = self.active_documents[document_id]["docx_object"]
            save_path = Path(file_path)
            
            # Ensure directory exists
            save_path.parent.mkdir(parents=True, exist_ok=True)
            
            if format.lower() == "docx":
                doc.save(str(save_path))
            elif format.lower() == "pdf":
                # PDF export would require AppleScript or additional libraries
                # For now, save as DOCX and note the limitation
                docx_path = save_path.with_suffix(".docx")
                doc.save(str(docx_path))
                logger.warning(f"PDF export not implemented, saved as DOCX: {docx_path}")
                save_path = docx_path
            else:
                raise ValueError(f"Unsupported format: {format}")
            
            # Update metadata
            self.active_documents[document_id]["metadata"]["file_path"] = str(save_path)
            
            logger.info(f"Saved document to {save_path}")
            return {
                "status": "success",
                "document_id": document_id,
                "file_path": str(save_path),
                "format": format
            }
            
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise
    
    async def _apply_paragraph_formatting(
        self,
        paragraph,
        formatting: Dict[str, Any]
    ) -> None:
        """Apply formatting to a paragraph.
        
        Args:
            paragraph: Word paragraph object
            formatting: Formatting options
        """
        try:
            # Alignment
            if "alignment" in formatting:
                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if formatting["alignment"] in align_map:
                    paragraph.alignment = align_map[formatting["alignment"]]
            
            # Apply formatting to runs (text formatting)
            if paragraph.runs:
                run = paragraph.runs[0]
                
                # Font size
                if "font_size" in formatting:
                    run.font.size = Pt(formatting["font_size"])
                
                # Font name
                if "font_name" in formatting:
                    run.font.name = formatting["font_name"]
                
                # Bold
                if "bold" in formatting:
                    run.font.bold = formatting["bold"]
                
                # Italic
                if "italic" in formatting:
                    run.font.italic = formatting["italic"]
                
                # Color (simplified - would need proper color handling)
                if "color" in formatting:
                    # python-docx color handling is more complex
                    # This is a simplified implementation
                    pass
            
        except Exception as e:
            logger.warning(f"Failed to apply some paragraph formatting: {e}")
    
    async def get_document_info(self, document_id: str) -> Dict[str, Any]:
        """Get information about a document.
        
        Args:
            document_id: ID of the document
            
        Returns:
            Dict with document information
        """
        if document_id not in self.active_documents:
            raise ValueError(f"Document {document_id} not found")
        
        return self.active_documents[document_id]["metadata"]
    
    async def list_documents(self) -> List[Dict[str, Any]]:
        """List all active documents.
        
        Returns:
            List of document metadata
        """
        return [data["metadata"] for data in self.active_documents.values()]
